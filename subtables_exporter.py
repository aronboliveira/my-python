import os
import sys
import platform
import time
import subprocess
import importlib
import re
from datetime import datetime
subt_acc = 0
failed_header = '#CABEÇALHO NÃO IDENTIFICADO'
starting_header_cell = ''
perm_error = f"Permission denied by {sys.platform} {platform.platform()} {platform.release()}. Recheck your permission."
doc = None

def prepare(package:str) -> bool:
    """
    Check if 'package' is importable. If not, prompt the user to install.
    Returns True if the package is importable (or installed successfully),
    otherwise False.
    """
    try:
      importlib.import_module(package)
      return True
    except:
        y = input(
            f'{package} is necessary and not installed in the system. '
            f'Prompt "Y" if you want to install it, else the script execution will be aborted:\n'
        ).strip().lower()
        if not y:
            return False
        try:
            subprocess.check_call([sys.executable, '-m', 'pip', 'install', package])
            return True
        except PermissionError:
            print(perm_error)
            sys.exit(1)
        except Exception as e:
            print('Error log:\n', e)
            print(f"Failed to install {package}. Exiting.")
            sys.exit(2)

def create_template(empresa: str, agendamentos: list[str]) -> str:
    agendamentos_str = "\n".join(agendamentos)
    t = f"""Olá, {empresa},

Obrigada por sua inscrição.
Com isso, a seguir estao seus horários de atendimento agendados de acordo com o interesse sinalizado no nosso formulário.

{agendamentos_str}

Aguardamos sua presença. 
Em caso de desistência, pedimos que nos sinalize por e-mail. 

Atenciosamente,
"""
    print(f"Olá, {empresa} ... {agendamentos_str} ...")
    return t
 
def process_subtable(subtable_df, start: int) -> list:
    try:
        global subt_acc
        global starting_header_cell
        subt_acc += 1
        print(Fore.BLUE + f"Processing subtable {subt_acc}...")
        header_idx = None
        print(Fore.BLUE + "Searching for headers...")
        for idx, row in subtable_df.iterrows():
            if any(re.match(r'^Empresa$', str(cell).strip(), re.IGNORECASE) for cell in row):
                header_idx = idx - start
                break
        if header_idx is None:
            print(Fore.YELLOW + 'ERROR: No header row found with "Empresa" match. Skipping subtable.')
            return []
        else:
            print(Fore.GREEN + f"Found header for identifying company: {header_idx + 1}")
        header_row = subtable_df.iloc[header_idx]
        print(Fore.BLUE + f"Header row for subtable {subt_acc}: " + str(header_row))
        start_col = 1 if (str(header_row[1]) and header_row[1] not in ['NaN', 'nan']) else 0
        if starting_header_cell and header_row[start_col] != starting_header_cell and (header_row[start_col] in ['nan', None, '']):
          for i, cell in enumerate(header_row):
            if str(cell).strip() and str(cell).strip().lower() != 'nan':
               start_col = i
               break
        if start_col is None:
            print(Fore.YELLOW + "WARNING: No non-empty header cell found; defaulting to column 0")
            start_col = 0
        headers = [str(cell).strip() if str(cell).strip().lower() != 'nan' 
                   else failed_header
                   for cell in header_row.iloc[start_col:]]
        if not headers or all(element == failed_header for element in headers):
            print(Fore.YELLOW + "ERROR: No valid headers found. Skipping subtable.")
            return []
        starting_header_cell = headers[0]
        print(Fore.GREEN + f"Found headers: {headers}")
        row0 = subtable_df.iloc[0].iloc[start_col:]
        heading_name = None
        for cell in row0:
            text = str(cell).strip()
            # PLACEHOLDER -> REPLACE WITH THE ACTUAL HEADING PATTERN
            if re.search(r'heading', text, re.IGNORECASE):
                heading_name = text
                break
        if not heading_name or heading_name.lower() == 'nan':
            print(Fore.YELLOW + "Failed to capture heading's name")
            return []
        print(Fore.BLUE + f"Reading data related to {heading_name}")
        if header_idx + 1 >= len(subtable_df):
            print(Fore.YELLOW + "ERROR: There's not enough rows after the header. Skipping subtable.")
            return []
        available_cols = subtable_df.shape[1] - start_col
        if available_cols < len(headers):
            print(Fore.YELLOW + "ERROR: Not enough columns in subtable. Skipping subtable.")
            return []
        data_df = subtable_df.iloc[header_idx + 1:, start_col:start_col + len(headers)]
        data_df.columns = headers
        csv_results = []
        xlsx_results = []
        print(Fore.BLUE + "Iterating over rows...")
        for _, row in data_df.iterrows():
            if row.isnull().all() or ('Empresa' in row and str(row['Empresa']).strip().lower() == 'intervalo'):
                continue
            row_dict = {header: str(row[header]).strip() for header in headers}
            row_dict['heading'] = heading_name
            print(Fore.RED + str(row_dict))
            failed_name = '#NOME NÃO ENCONTRADO'
            empresa = next((row_dict.get(key, '').strip() for key in ('Empresa','empresa') 
                            if row_dict.get(key, '').strip()), failed_name)
            if empresa == failed_name:
                print(Fore.YELLOW + "WARN: Failed to get company name")
            if not empresa or empresa == failed_name:
                print(Fore.YELLOW + "WARN: Company name looks like a time. Falling back to 'Nome' column.")
                empresa = "Representante" + next((row_dict.get(key, '').strip() for key in ('Nome','nome')
                                if row_dict.get(key, '').strip()), failed_name).strip()
            row_dict['empresa'] = empresa
            failed_timestamp = '#HORARIO INDEFINIDO'
            hora = next((row_dict.get(k, '').strip() for k in ('Hora','hora','Horário','horário','Horario','horario')
                         if row_dict.get(k, '').strip()), failed_timestamp)
            if hora == failed_timestamp:
                print(Fore.YELLOW + "WARN: Failed to get timestamp for the invitation")
            row_dict['hora'] = hora
            print(Fore.BLUE + f"Finished selecting data for writing email to {row_dict['empresa']}. Starting template...")
            for k in row_dict:
                if str(row_dict[k]).strip().lower() == 'nan':
                    row_dict[k] = 'Não preenchido'
            csv_results.append(row_dict)
        print(Fore.GREEN + f"Finished processing subtable {subt_acc} successfully")
        return csv_results
    except Exception as e:
        print(Fore.RED + f"Error processing subtable: {str(e)}")
        return []

def main():
    tries = 0
    max_tries = 10
    global doc
    doc = Document()
    doc_title = doc.add_heading('Agenda de Cabeçalhos de Atendimento', level=1)
    doc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph().add_run('\n\n')
    while tries < max_tries:
        tries += 1
        print(Fore.CYAN + f"\nAttempt {tries}/{max_tries}")
        file_path = os.path.relpath(input(Style.BRIGHT + "Enter the relative file path (or exit the terminal with Ctrl + C), including the extension: ")).strip()
        if not os.path.isfile(file_path):
            print(Fore.RED + "File not found. Please try again.")
            continue
        ext = os.path.splitext(file_path)[1].lower()
        if ext not in ['.csv', '.xlsx', '.xls']:
            print(Fore.RED + "Invalid file format. Only CSV, XLSX, and XLS are supported.")
            continue
        try:
            print(Fore.BLUE + "Converting to Dataframe...")
            if ext == '.csv':
                df = pd.read_csv(file_path, header=None, dtype=str)
                print (Fore.GREEN + "Successfully created dataframe using the .csv!")
            else:
                df = pd.read_excel(file_path, header=None, engine='openpyxl', dtype=str)
                print (Fore.GREEN + "Successfully created dataframe using the spreadsheet file!")
            heading_indexes = []
            pattern = re.compile(r'^balc[aã]o', re.IGNORECASE)
            print(Fore.BLUE + 'Looking for "Heading" headers...')
            for idx, row in df.iterrows():
                for cell in row:
                    if pd.isnull(cell): continue
                    cell_value = str(cell).strip().lower()
                    if pattern.match(cell_value):
                        heading_indexes.append(idx)
                        print(Fore.GREEN + f"Found Heading {len(heading_indexes)}")
                        break
            if not heading_indexes:
                print(Fore.YELLOW + 'No "Heading" sections were found in the file.')
                return
            all_csv_data = []
            all_xlsx_data = []
            print(Fore.BLUE + "Going through indexes for the Cabeçalhos...")
            for i in range(len(heading_indexes)):
                start = heading_indexes[i]
                end = heading_indexes[i+1] if i+1 < len(heading_indexes) else len(df)
                print(Fore.BLUE + f"Working heading table starting at row {start + 1}, ranging to {end + 1}")
                subtable = df.iloc[start:end]
                if subt_acc > 1:
                  print(Fore.YELLOW + "Giving you 3 seconds to read! ☺")
                  time.sleep(3)
                  print(Fore.YELLOW + "Continuing...")
                res = process_subtable(subtable, start)
                print(Fore.CYAN + "Result from processing: " + str([{k: v for k, v in d.items() if k != 'email_template'} for d in res]))
                all_csv_data.extend(res)
            if not all_csv_data:
                print(Fore.YELLOW + "No valid data found to process.")
                continue
            company_schedules = {}
            for d in all_csv_data:
                empresa = d['empresa'].strip()
                hora = d['hora']
                heading = d['heading']
                if empresa not in company_schedules:
                    company_schedules[empresa] = []
                company_schedules[empresa].append(f"{hora} — {heading}")
            for empresa, agendamentos in sorted(company_schedules.items(), key=lambda item: item[0].strip()):
                all_xlsx_data.append({"Empresa": empresa, "Corpo do email": create_template(empresa=empresa, agendamentos=agendamentos)})
                agendamentos.sort()
                run = doc.add_paragraph().add_run(empresa)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 102, 204)
                run.font.size = Pt(14)
                for a in agendamentos:
                    doc.add_paragraph(a).style.font.size = Pt(12)
                doc.add_paragraph().add_run('\n')
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_csv = f"locais_de_atendimento_{timestamp}.csv"
            print(Fore.CYAN + f"Outputting to path {output_csv}")
            pd.DataFrame(all_csv_data).to_csv(output_csv, index=False, encoding='utf-8-sig')
            output_xlsx = f"templates_de_mensagem_{timestamp}.xlsx"
            print(Fore.CYAN + f"Outputting to path {output_xlsx}")
            pd.DataFrame(all_xlsx_data).to_excel(output_xlsx, index=False, engine='openpyxl')
            wb = load_workbook(output_xlsx)
            ws = wb.active
            even_fill = PatternFill(start_color="00FFFFFF", end_color="00FFFFFF", fill_type="solid")
            odd_fill = PatternFill(start_color="FFFFFDD0", end_color="FFFFFDD0", fill_type="solid")
            for col_cells in ws.columns:
                max_length = 0
                column = col_cells[0].column_letter
                for cell in col_cells:
                    if cell.value:
                        length = len(str(cell.value))
                        max_length = max(max_length, length)
                adjusted_width = min(max_length + 2, 30)
                ws.column_dimensions[column].width = adjusted_width
            for col in range(1, ws.max_column + 1):
                cell = ws.cell(row=1, column=col)
                cell.fill = PatternFill(start_color="FF00FF00", end_color="FF00FF00", fill_type="solid")
                cell.font = Font(bold=True)
            for row in range(2, ws.max_row + 1):
                fill = even_fill if row % 2 == 0 else odd_fill
                for col in range(1, ws.max_column + 1):
                    ws.cell(row=row, column=col).fill = fill
            wb.save(output_xlsx)
            output_docx = f"relacao_de_horarios_{timestamp}.docx"
            print(Fore.CYAN + f"Outputting to path {output_docx}")
            doc.save(output_docx)
            print(Fore.GREEN + f"\nSuccess! Output files created as: ({output_csv}, {output_xlsx})")
            return
        except PermissionError as pe:
            print(Fore.YELLOW + PermissionError)
            print(Fore.YELLOW + f"Try closing the file or checking permissions")
        except Exception as e:
            print(Fore.RED + f"Error processing file: {str(e)}")
    print(Fore.RED + "\nMaximum attempts reached. Exiting.")

if __name__ == "__main__":
    all_prepared = []
    for package in ["pandas", "openpyxl", "colorama", "python-docx"]:
      all_prepared.append(prepare(package))
    if False in all_prepared:
        print("Could not proceed with execution due to missing dependencies or an undefined error. Exiting")
        sys.exit(3)
    from colorama import Fore, init, Style
    import pandas as pd
    from openpyxl import load_workbook
    from openpyxl.styles import Font, PatternFill
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    init(autoreset=True)
    main()
